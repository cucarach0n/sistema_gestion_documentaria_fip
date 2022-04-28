# -*- coding: utf-8 -*-
import os
from apps.file.models import File, FileInFolder
from apps.folder.models import Folder
from rest_framework.response import Response
from rest_framework import status
from apps.file.api.serializers.file_serializers import (
    FileCreateSerializer, FileHistorySerializer,FileObtenerSerializer,FileFolderCreateSerializer,
    FileDetalleSerializer,FileUpdateOcrSerializer,FileBuscarSerializer, FileUpdatePrivateSerializer, FileUpdateSerializer
    )
from apps.file.api.serializers.general_serializers import File_Serializer,FileInFolder_Serializer
from apps.users.authenticacion_mixings import Authentication
from django.core.files.storage import FileSystemStorage
from django.conf import settings
from rest_framework import viewsets
from django.http import FileResponse
from apps.base.util import DocumentoOCR, createHistory, setHistory, validarPrivado
import threading
import PyPDF2
from PyPDF2 import PdfFileReader, PdfFileMerger,PdfFileWriter
import pathlib
from docx import Document
#import pdfplumber
#from apps.base.pdfConvertPdfMiner import extrarText
from django.utils.crypto import get_random_string
#import fitz

from django.db.models import Q
from django.core.paginator import Paginator

def guardarOcr(file,id,idUser,textoExtraido):
    '''if textoExtraido == "":
        documentoRenderisado = DocumentoOCR(file)
        text = str(documentoRenderisado.obtenerTexto())
              
    elif textoExtraido:
        text = textoExtraido'''
    documentoRenderisado = DocumentoOCR(file)
    text = str(documentoRenderisado.obtenerTexto())
    file = FileUpdateOcrSerializer(File.objects.filter(id = id).first(),data = {'contenidoOCR':text})
    
    if file.is_valid():  
        fileSave = file.save()
        #set history file
        setHistory(fileSave,'contenido OCR registrado',idUser)
def extraerTextDocx(pathFile,idFile):
    doc = Document(pathFile)
    text =""
    for paragraph in doc.paragraphs:
        text += paragraph.text
    File.objects.filter(id = idFile).update(contenidoOCR = text)
def extraerExtencion(Archivo):
    '''extension = [["jpg","image/jpg"]
                ,["jpeg","image/jpeg"]
                ,["png","image/png"]
                ,["gif","image/gif"]
                ,["xlsx","application/vnd.ms-excel"]
                ,["docx","application/msword"]
                ,["pptx","application/vnd.ms-powerpoint"]
                ,["pdf","application/pdf"]
                ,["txt","text/plain;charset=UTF-8"]
                ,["zip","application/zip"]
                ,["rar","application/x-rar-compressed"]
                ,["mp4","audio/mp4"]
                ,["mpeg","video/mpeg"]
                ]'''
    extension = [[".jpg","image/jpeg"]
                ,[".jpeg","image/jpeg"]
                ,[".png","image/png"]
                ,[".gif","image/gif"]
                ,[".pdf","application/pdf"]
                ,[".txt","text/plain"]
                ,[".mp4","video/mp4"]
                ,[".mpeg","video/mpeg"]
                ,[".xlsx","application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]
                ,[".xls","application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"]
                ,[".docx","application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
                ,[".doc","application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
                ,[".pptx","application/vnd.openxmlformats-officedocument.presentationml.presentation"]
                ,[".zip","application/x-zip-compressed"]
                ,[".rar","application/x-rar-compressed"]
                ]
    path = pathlib.Path(settings.MEDIA_ROOT+'files/' + Archivo)
    ext,aplication = None,None
    for e in extension:
        if e[0] in ''.join(path.suffixes):
            ext = e[0][1:]
            aplication =e[1]
            return ext,aplication      
    if ext == None:
        ext =''.join(path.suffixes)[1:]
        aplication = ''
        return ext,aplication
def normalisarNameDocument(nameFile):
    caracteres = "%()$&#~`+^*=,;°"
    #caracteres = "!#$%^&*()"
    for caracteres in caracteres:
        nameFile = nameFile.replace(caracteres,'')
    nameFile = nameFile.replace(" ","_")
    return nameFile
        
def saveFile(self,request,documento_serializer,scope):
    #current_site = get_current_site(request).domain
    ruta = settings.MEDIA_ROOT+'files/'
    fs = FileSystemStorage(location=ruta)
    
    #nameFile = request.FILES['documento_file'].name.replace(" ","_")
    nameFile = normalisarNameDocument(request.FILES['documento_file'].name)
    
    file_in = request.FILES['documento_file']

    pdf_merger = PdfFileMerger()
    pdfreadRewrite = PdfFileReader(file_in,  strict = False)
    pdfwrite = PdfFileWriter()
    for page_count in range(pdfreadRewrite.numPages):
        pages = pdfreadRewrite.getPage(page_count)
        pdfwrite.addPage(pages)

    fileobjfix = open(settings.MEDIA_ROOT+'test/'+'fixedPDF.pdf', 'w+b')
    pdfwrite.write(fileobjfix)
    #fileobjfix.close()

    pdf_merger.append(fileobjfix)
    pdf_merger.addMetadata({
        '/Title': str(documento_serializer.validated_data['nombreDocumento'])
    })
    file_out = open(settings.MEDIA_ROOT+'test/'+'new.pdf', 'w+b')
    pdf_merger.write(file_out)

    '''file_in.close()
    file_out.close()'''
    
    #file = fs.save(nameFile,request.FILES['documento_file'])
    file = fs.save(nameFile,file_out)
    fileurl = fs.get_valid_name(file)
    #documento_serializer.validated_data['documento_file'] = doc
    #documento_serializer.validated_data['contenidoOCR'] = "test"
    #documento = documento_serializer.save()
    #set history unidadArea
    #setHistory(documento,'registro file',self.userFull.id)
    documento = File()
    documento.nombreDocumento = documento_serializer.validated_data['nombreDocumento']
    documento.documento_file = fileurl
    documento.user_id = self.userFull.id
    documento.scope = scope
    #documento.contenidoOCR = "-"
    documento.extension,application = extraerExtencion(fileurl)
    documento.slug = get_random_string(length=11)
    documento.unidadArea_id = self.userFull.unidadArea_id
    #mensaje = "Documento cargado exitosamente"
    #documentoOcr = DocumentoOCR(fileurl)
    textPDF = '' #obtenerTextoPDF("/"+fileurl)
    #documento.contenidoOCR = textPDF #documentoOcr.obtenerTexto()
    documento.save()
    #set history unidadArea
    setHistory(documento,'se registro el file',self.userFull.id)
    if documento.extension == "pdf":
        #if textPDF == "":
        #mensaje = 'Documento cargado exitosamente, se estra procesando el contenido del archivo...'                     
        threading_text = threading.Thread(target=guardarOcr,args=("/"+fileurl,documento.id,self.userFull.id,textPDF,))
        threading_text.start()
        #print('Cantidad de threading : ',threading.active_count())
    if documento.extension == "docx":
        #extraerTextDocx(ruta + fileurl,documento.id)
        path = ruta + fileurl
        threading_docx = threading.Thread(target=extraerTextDocx,args=(path,documento.id,))
        threading_docx.start()
        '''doc = Document(ruta + fileurl)
        for paragraph in doc.paragraphs:
            print(paragraph.text)'''
    #print(documento)
    fileDetalleSerializer = FileDetalleSerializer(documento,context = {'userId':self.userFull.id})
    #print(fileDetalleSerializer)
    #set history file
    #setHistory(documento,'agrego datos del file',self.userFull.id)
    parentID = Folder.objects.filter(slug = request.data['directorioslug']).first()
    if parentID:
        fileinfoler_serializer = FileInFolder_Serializer(data = {
            'file': documento.id,
            'parent_folder':parentID.id#Folder.objects.filter(slug = request.data['directorioslug']).first().id
        })
        if fileinfoler_serializer.is_valid():
            fileInFolderSave = fileinfoler_serializer.save()
            #set history fileinfolder
            setHistory(fileInFolderSave,'registro en el folder',self.userFull.id)
            #add history folder
            createHistory(Folder,parentID.id,"Se agrego el file " + documento.documento_file.name,"+",self.userFull.id)
            '''history = Folder.historical.create(id=parentID.id,history_date = datetime.today()
                                        ,history_change_reason = "Se agrego el file " + documento.documento_file.name
                                        ,history_type = "+", history_user_id = self.userFull.id )
            history.save()'''
        '''threading_text = threading.Thread(target=guardarOcr,args=(fileurl,documento.id,))
        threading_text.start()
        print('Cantidad de threading : ',threading.active_count())'''
        return Response(fileDetalleSerializer.data,status = status.HTTP_200_OK)
    #File.objects.filter(id = documento.id).delete()
    return Response({'error':'La carpeta contenedora no existe'},status = status.HTTP_404_NOT_FOUND)   
class FileObtenerViewSet(Authentication,viewsets.GenericViewSet):
    serializer_class = FileObtenerSerializer
    def get_queryset(self,pk=None):
        if self.userFull.is_staff < 3:
            return self.serializer_class().Meta.model.objects.filter(slug=pk,unidadArea_id = self.userFull.unidadArea_id,eliminado = False).first()
        return self.serializer_class().Meta.model.objects.filter(slug=pk,eliminado = False).first()
    def retrieve(self,request,pk=None):

        #documento_query = self.get_queryset(pk)
        #if documento_query:
        documento = self.get_queryset(pk)
        
        '''if not(documento_query.scope == False and documento_query.user_id == self.userFull.id):
            return Response({'error':'El archivo solicitado es privado'},status = status.HTTP_401_UNAUTHORIZED)'''
        #if documento_query:
        #documento = File_Serializer().Meta.model.objects.filter(slug=pk,unidadArea_id = self.userFull.unidadArea_id).first()
        if documento:
            if validarPrivado(documento,self.userFull.id,True):
                return Response({'error':'La carpeta contenedora o el file son privados'},status = status.HTTP_401_UNAUTHORIZED)
            #test copy file####
            '''rutaFile = settings.MEDIA_ROOT+'files/'
            fileObject = open(rutaFile+documento.documento_file.name, 'rb')
            fs = FileSystemStorage(location=rutaFile)
            file = fs.save(documento.documento_file.name,fileObject)
            nameNewFile = fs.get_valid_name(file)
            print(nameNewFile)'''
            ###################
            print('Sirviendo file {0} al usuario {1}'.format(documento.documento_file.name,self.userFull.correo))
            doc = str(documento.documento_file)#56 en linux / 34 windows
            file_location =  settings.MEDIA_ROOT + 'files/' + doc 
            try:    
                #with open(file_location, 'r') as f:
                #    file_data = f.read()
                file_data = open(file_location, 'rb')
                # sending response 
                ext,app = extraerExtencion(str(documento.documento_file))#56 en linux/ 34 windows
                
                response = FileResponse(file_data, content_type=app)
            
                #response['Content-Length'] = file_data.size
                response['Content-Disposition'] = 'attachment; filename="'+str(documento.nombreDocumento)+'"'#56 en linux/ 34 windows
                #create history
                createHistory(File,documento.id,"Documento " + documento.documento_file.name + " visto","v",self.userFull.id)
            except:
                # handle file not exist case here
                return Response({'error':'Hubo un error al obtener el archivo'},status = status.HTTP_400_BAD_REQUEST)
            return response
        return Response({'error':'No existe el documento o archivo solicitado'},status = status.HTTP_404_NOT_FOUND)
        
        #return Response({'error':'Error al procesar la solicitud'},status = status.HTTP_400_BAD_REQUEST)
        #return Response({'error':'No existe el file'},status = status.HTTP_400_BAD_REQUEST)
class FileListViewSet(Authentication,viewsets.GenericViewSet):
    serializer_class = FileUpdateSerializer

    def get_queryset(self,pk=None):
        if pk is None:
            return FileDetalleSerializer.Meta.model.objects.filter(Q(user_id= self.userFull.id,scope = False,unidadArea_id = self.userFull.unidadArea_id,eliminado = False)
                                                                |Q(scope = True,unidadArea_id = self.userFull.unidadArea_id,eliminado = False)).distinct()
        if self.userFull.is_staff < 3:
            return FileDetalleSerializer.Meta.model.objects.filter(slug=pk,unidadArea_id = self.userFull.unidadArea_id,eliminado = False).first()
        return FileDetalleSerializer.Meta.model.objects.filter(slug=pk,eliminado = False).first()
    def list(self,request):
        documento_serializer = FileDetalleSerializer(self.get_queryset(),many = True,context = {'userId':self.userFull.id})
        return Response(documento_serializer.data,status = status.HTTP_200_OK)
    def retrieve(self,request,pk=None):
        documentoResult = self.get_queryset(pk)
        if documentoResult:
            documentoSerializer = FileDetalleSerializer(self.get_queryset(pk),context = {'userId':self.userFull.id})
            docResult = self.get_queryset(pk)
            #if not docResult.scope:
            #    return Response({'error':'Este file no es publico'},status = status.HTTP_401_UNAUTHORIZED)
            if validarPrivado(self.get_queryset(pk),self.userFull.id,True):
                return Response({'error':'La carpeta contenedora o el file son privados'},status = status.HTTP_401_UNAUTHORIZED)
            '''if not(documento.scope == False and documento.user_id == self.userFull.id):
                return Response(documento.data,status=status.HTTP_200_OK)
            return Response({'error':'El archivo solicitado es privado'},status = status.HTTP_401_UNAUTHORIZED)'''
            #create history
            createHistory(File,docResult.id,"Obteniendo documento " + docResult.documento_file.name,"o",self.userFull.id)
            return Response(documentoSerializer.data,status=status.HTTP_200_OK)
        return Response({'error':'No existe el documento solicitado'},status = status.HTTP_404_NOT_FOUND)
    def update(self,request,pk=None):
        documento = self.get_queryset(pk)
        if documento:
            #if not documento.scope:
            #    return Response({'error':'Este file no es publico'},status = status.HTTP_401_UNAUTHORIZED)
            if validarPrivado(documento,self.userFull.id,True):
                return Response({'error':'La carpeta contenedora o el file son privados'},status = status.HTTP_401_UNAUTHORIZED)
            documento_serializer = self.get_serializer(documento,data = request.data)
            if documento_serializer.is_valid():
                #if bool(documento_serializer.validated_data['scope']) == False:
                #    documento_serializer.validated_data['user_id'] = self.userFull.id
                fileUpdate = documento_serializer.save()
                #set history file
                setHistory(fileUpdate,'actualizo file',self.userFull.id)
                return Response({'mensaje':'Documento actualizado correctamente'},status = status.HTTP_200_OK) 
            return Response({'error':'hubo un error al actualizar los datos'},status = status.HTTP_400_BAD_REQUEST)
        return Response({'error':'No existe el documento'},status = status.HTTP_404_NOT_FOUND)
class FileUpdatePrivateViewSet(Authentication,viewsets.GenericViewSet):
    serializer_class = FileUpdatePrivateSerializer

    def get_queryset(self,pk=None):
        if pk is None:
            return FileDetalleSerializer.Meta.model.objects.filter(Q(user_id= self.userFull.id,scope = False,unidadArea_id = self.userFull.unidadArea_id,eliminado = False)
                                                                |Q(scope = True,unidadArea_id = self.userFull.unidadArea_id,eliminado = False)).distinct()
        return FileDetalleSerializer.Meta.model.objects.filter(slug=pk,unidadArea_id = self.userFull.unidadArea_id,eliminado = False).first()
    def update(self,request,pk=None):
        documento = self.get_queryset(pk)
        if documento:
            #if not documento.scope:
            #    return Response({'error':'Este file no es publico'},status = status.HTTP_401_UNAUTHORIZED)
            if validarPrivado(documento,self.userFull.id,True):
                return Response({'error':'La carpeta contenedora o el file son privados'},status = status.HTTP_401_UNAUTHORIZED)
            documento_serializer = self.get_serializer(documento,data = request.data)
            if documento_serializer.is_valid():
                if bool(documento_serializer.validated_data['scope']) == True:
                    folderGestion = Folder.objects.filter(carpeta_hija__isnull =True,unidadArea_id = self.userFull.unidadArea_id).first()
                    FileInFolder.objects.filter(file_id = documento.id).update(parent_folder_id = folderGestion.id)
                    documento_serializer.validated_data['user_id'] = self.userFull.id
                    createHistory(File,documento.id,"Cambiando a publico "+ documento.nombreDocumento,"P",self.userFull.id)
                fileUpdate = documento_serializer.save()
                #set history file
                setHistory(fileUpdate,'actualizo file',self.userFull.id)
                return Response({'mensaje':'Documento actualizado correctamente'},status = status.HTTP_200_OK) 
            return Response({'error':'hubo un error al actualizar los datos'},status = status.HTTP_400_BAD_REQUEST)
        return Response({'error':'No existe el documento'},status = status.HTTP_404_NOT_FOUND)

class FileViewSet(Authentication,viewsets.GenericViewSet):
    
    serializer_class = FileFolderCreateSerializer
    '''
    def get_queryset(self,pk=None):
        if pk is None:
            return FileCreateSerializer().Meta.model.objects.filter(unidadArea_id = self.userFull.unidadArea_id)
        return FileCreateSerializer().Meta.model.objects.filter(slug=pk,unidadArea_id = self.userFull.unidadArea_id).first()
    def list(self,request):

        print(self.userFull.unidadArea_id)
        documento_serializer = FileDetalleSerializer(self.get_queryset(),many = True)
        return Response(documento_serializer.data,status = status.HTTP_200_OK)
    '''
    
    #post
    def create(self,request):
        documento_serializer = self.serializer_class(data = request.data,context ={"folderSlug":request.data['directorioslug']})
        
        if documento_serializer.is_valid():
            folderResult = Folder.objects.filter(slug = documento_serializer.validated_data['directorioslug'],unidadArea_id = self.userFull.unidadArea_id,eliminado = False)
            '''if not folderResult.first().scope:
                return Response({'error':'No puede crear un file publico aqui'},status = status.HTTP_401_UNAUTHORIZED)'''
            if folderResult:
                if validarPrivado(folderResult.first(),self.userFull.id):
                    return Response({'error':'La carpeta solicitada es privada'},status = status.HTTP_401_UNAUTHORIZED)
                '''padrePrivate = obtenerRuta(folderResult.first().id,[folderResult.first().nombre],True,False,True,self.userFull.id)
                if padrePrivate:
                    return Response({'error':'La carpeta es privada'},status = status.HTTP_401_UNAUTHORIZED)
                elif folderResult.first().scope == False:
                    if not(folderResult.first().user_id == self.userFull.id):
                        return Response({'error':'La carpeta es privada'},status = status.HTTP_401_UNAUTHORIZED)'''
                if folderResult:

                    '''#current_site = get_current_site(request).domain
                    ruta = settings.MEDIA_ROOT+'files/'
                    fs = FileSystemStorage(location=ruta)
                    nameFile = request.FILES['documento_file'].name.replace(" ","_")
                    print(nameFile)
                    file = fs.save(nameFile,request.FILES['documento_file'])
                    fileurl = fs.url(file)
                    #documento_serializer.validated_data['documento_file'] = doc
                    #documento_serializer.validated_data['contenidoOCR'] = "test"
                    #documento = documento_serializer.save()
                    #set history unidadArea
                    #setHistory(documento,'registro file',self.userFull.id)
                    documento = File()
                    documento.nombreDocumento = documento_serializer.validated_data['nombreDocumento']
                    documento.documento_file = nameFile
                    documento.user_id = self.userFull.id
                    documento.scope = True
                    documento.extension,application = extraerExtencion(fileurl[1:])
                    documento.slug = get_random_string(length=11)
                    documento.unidadArea_id = self.userFull.unidadArea_id
                    #mensaje = "Documento cargado exitosamente"
                    #documentoOcr = DocumentoOCR(fileurl)
                    textPDF = obtenerTextoPDF(fileurl)
                    #documento.contenidoOCR = textPDF #documentoOcr.obtenerTexto()
                    documento.save()
                    #set history unidadArea
                    setHistory(documento,'se registro el file',self.userFull.id)
                    if documento.extension == "pdf":
                        #if textPDF == "":
                        #mensaje = 'Documento cargado exitosamente, se estra procesando el contenido del archivo...'                     
                        threading_text = threading.Thread(target=guardarOcr,args=(fileurl,documento.id,self.userFull.id,textPDF,))
                        threading_text.start()
                        #print('Cantidad de threading : ',threading.active_count())
                    
                    #print(documento)
                    fileDetalleSerializer = FileDetalleSerializer(documento)
                    #print(fileDetalleSerializer)
                    #set history file
                    #setHistory(documento,'agrego datos del file',self.userFull.id)
                    parentID = Folder.objects.filter(slug = request.data['directorioslug']).first()
                    if parentID:
                        fileinfoler_serializer = FileInFolder_Serializer(data = {
                            'file': documento.id,
                            'parent_folder':parentID.id#Folder.objects.filter(slug = request.data['directorioslug']).first().id
                        })
                        if fileinfoler_serializer.is_valid():
                            fileInFolderSave = fileinfoler_serializer.save()
                            #set history fileinfolder
                            setHistory(fileInFolderSave,'registro en el folder',self.userFull.id)
                            #add history folder
                            createHistory(Folder,parentID.id,"Se agrego el file " + documento.documento_file.name,"+",self.userFull.id)
                        return Response(fileDetalleSerializer.data,status = status.HTTP_200_OK)
                    #File.objects.filter(id = documento.id).delete()
                    return Response({'error':'La carpeta contenedora no existe'},status = status.HTTP_404_NOT_FOUND)'''
                    return saveFile(self,request,documento_serializer,folderResult.first().scope) 
                return Response({'error':'La carpeta contenedora no es accesible para su usuario'},status = status.HTTP_401_UNAUTHORIZED)  
            return Response({'error':'La carpeta contenedora no existe'},status = status.HTTP_401_UNAUTHORIZED)   
        else:
            #return Response({'Error':'no se pudo cargar el documento'},status = status.HTTP_400_BAD_REQUEST)
            return Response(documento_serializer.errors,status = status.HTTP_400_BAD_REQUEST)
    '''def retrieve(self,request,pk=None):
        documento = FileDetalleSerializer(self.get_queryset(pk))
        if documento:
            return Response(documento.data,status=status.HTTP_200_OK)
        return Response({'error':'No existe el documento solicitado'},status = status.HTTP_404_NOT_FOUND)'''
    '''def update(self,request,pk=None):
        documento = self.get_queryset(pk)
        if documento:
            documento_serializer = self.serializer_class(documento,data = request.data)
            if documento_serializer.is_valid():
                documento_serializer.save()
                return Response({'mensaje':'Documento actualizado correctamente'},status = status.HTTP_200_OK)                
            return Response({'error':'hubo un error al actualizar los datos'},status = status.HTTP_400_BAD_REQUEST)
        return Response({'error':'No existe el documento'},status = status.HTTP_400_BAD_REQUEST)'''
#pendiente eliminar
class FilePrivateViewSet(Authentication,viewsets.GenericViewSet):
    serializer_class = FileFolderCreateSerializer
    def create(self,request):
        documento_serializer = self.serializer_class(data = request.data)
        if documento_serializer.is_valid():
            folderResult = Folder.objects.filter(slug = documento_serializer.validated_data['directorioslug'],unidadArea_id = self.userFull.unidadArea_id,eliminado = False)
            '''if folderResult.first().scope:
                return Response({'error':'No puede crear el file aqui'},status = status.HTTP_401_UNAUTHORIZED)'''
            if folderResult.first().scope:
                #if not Folder.objects.filter(carpeta_hija__isnull =True,unidadArea_id = self.userFull.unidadArea_id,slug = folderResult.first().slug).first():
                return Response({'error':'No puede crear un file privado aqui'},status = status.HTTP_401_UNAUTHORIZED)
            if validarPrivado(folderResult.first(),self.userFull.id):
                return Response({'error':'La carpeta solicitada es privada'},status = status.HTTP_401_UNAUTHORIZED)
            if folderResult:
                return saveFile(self,request,documento_serializer,False)
            return Response({'error':'La carpeta contenedora no es accesible para su usuario'},status = status.HTTP_401_UNAUTHORIZED)
        else:
            return Response(documento_serializer.errors,status = status.HTTP_400_BAD_REQUEST)

class FileBuscarAPIView(Authentication,viewsets.GenericViewSet):

    serializer_class = FileBuscarSerializer
    def get_queryset(self,data):
        '''if data['opcion'] == 1:
            return self.get_serializer().Meta.model.objects.filter(Q(nombreDocumento__icontains = data['buscar'],scope =False,user_id=self.userFull.id,eliminado = False)| 
                                                                Q(contenidoOCR__icontains = data['buscar'],scope =False,user_id=self.userFull.id,eliminado = False)|
                                                                Q(nombreDocumento__icontains = data['buscar'],scope =True,eliminado = False)| 
                                                                Q(contenidoOCR__icontains = data['buscar'],scope =True,eliminado = False)
                                                                ,unidadArea_id = self.userFull.unidadArea_id,eliminado = False).distinct()
        elif data['opcion'] == 2:
            return self.get_serializer().Meta.model.objects.filter(Q(filetag__tag__tagName__icontains = data['buscar'],
                                                                    unidadArea_id = self.userFull.unidadArea_id,scope =False,user_id=self.userFull.id,eliminado = False)|
                                                                    Q(filetag__tag__tagName__icontains = data['buscar'],
                                                                    unidadArea_id = self.userFull.unidadArea_id,scope =True,eliminado = False))

        elif data['opcion'] == 3:
            return self.get_serializer().Meta.model.objects.filter(Q(etiqueta__nombre__icontains = data['buscar'],
                                                                    unidadArea_id = self.userFull.unidadArea_id,
                                                                    scope =False,user_id=self.userFull.id,eliminado = False)|
                                                                    Q(etiqueta__nombre__icontains = data['buscar'],
                                                                    unidadArea_id = self.userFull.unidadArea_id,
                                                                    scope =True,eliminado = False))
        elif data['opcion'] == 4:
            return self.get_serializer().Meta.model.objects.filter(Q(nombreDocumento__icontains = data['buscar'],scope =False,user_id=self.userFull.id,eliminado = False)| 
                                                                Q(contenidoOCR__icontains = data['buscar'],scope =False,user_id=self.userFull.id,eliminado = False)| 
                                                                Q(filetag__tag__tagName__icontains = data['buscar'],scope =False,user_id=self.userFull.id,eliminado = False)| 
                                                                Q(etiqueta__nombre__icontains = data['buscar'],scope =False,user_id=self.userFull.id,eliminado = False)|

                                                                Q(nombreDocumento__icontains = data['buscar'],scope =True,eliminado = False)| 
                                                                Q(contenidoOCR__icontains = data['buscar'],scope =True,eliminado = False)| 
                                                                Q(filetag__tag__tagName__icontains = data['buscar'],scope =True,eliminado = False)| 
                                                                Q(etiqueta__nombre__icontains = data['buscar'],scope =True,eliminado = False)
                                                                ,unidadArea_id = self.userFull.unidadArea_id,eliminado = False).distinct()'''
        if self.userFull.is_staff < 3:
            if data['opcion'] == 1:
                return self.get_serializer().Meta.model.objects.filter(Q(nombreDocumento__icontains = data['buscar'])| 
                                                                    Q(contenidoOCR__icontains = data['buscar'])
                                                                    ,Q(scope = True)|Q(scope = False, user_id = self.userFull.id),
                                                                    unidadArea_id = self.userFull.unidadArea_id,eliminado = False).distinct()
            elif data['opcion'] == 2:
                return self.get_serializer().Meta.model.objects.filter(Q(filetag__tag__tagName__icontains = data['buscar'])
                                                                    ,Q(scope = True)|Q(scope = False, user_id = self.userFull.id),
                                                                    unidadArea_id = self.userFull.unidadArea_id,eliminado = False).distinct()

            elif data['opcion'] == 3:
                return self.get_serializer().Meta.model.objects.filter(Q(etiqueta__nombre__icontains = data['buscar'])
                                                                    ,Q(scope = True)|Q(scope = False, user_id = self.userFull.id),
                                                                    unidadArea_id = self.userFull.unidadArea_id,eliminado = False).distinct()
            elif data['opcion'] == 4:
                return self.get_serializer().Meta.model.objects.filter(Q(nombreDocumento__icontains = data['buscar'])| 
                                                                    Q(contenidoOCR__icontains = data['buscar'])| 
                                                                    Q(filetag__tag__tagName__icontains = data['buscar'])| 
                                                                    Q(etiqueta__nombre__icontains = data['buscar'])
                                                                    ,Q(scope = True)|Q(scope = False, user_id = self.userFull.id),
                                                                    unidadArea_id = self.userFull.unidadArea_id,eliminado = False).distinct()
        
        if data['opcion'] == 1:
            return self.get_serializer().Meta.model.objects.filter(Q(nombreDocumento__icontains = data['buscar'])| 
                                                                Q(contenidoOCR__icontains = data['buscar'])
                                                                ,scope = True,eliminado = False).distinct()
        elif data['opcion'] == 2:
            return self.get_serializer().Meta.model.objects.filter(Q(filetag__tag__tagName__icontains = data['buscar'])
                                                                ,scope = True,eliminado = False).distinct()

        elif data['opcion'] == 3:
            return self.get_serializer().Meta.model.objects.filter(Q(etiqueta__nombre__icontains = data['buscar'])
                                                                ,scope = True,eliminado = False).distinc()
        elif data['opcion'] == 4:
            return self.get_serializer().Meta.model.objects.filter(Q(nombreDocumento__icontains = data['buscar'])| 
                                                                Q(contenidoOCR__icontains = data['buscar'])| 
                                                                Q(filetag__tag__tagName__icontains = data['buscar'])| 
                                                                Q(etiqueta__nombre__icontains = data['buscar'])
                                                                ,scope = True,eliminado = False).distinct()
    def create(self,request):
            file_serializer = self.get_serializer(data = request.data)
            if file_serializer.is_valid():
                fileBusqueda_serializer = FileDetalleSerializer(self.get_queryset(file_serializer.data),many = True,context = {'userId':self.userFull.id})
                return Response(fileBusqueda_serializer.data,status = status.HTTP_200_OK)
            else:
                return Response(file_serializer.errors,status = status.HTTP_400_BAD_REQUEST)
class FileHistoryAPIView(Authentication,viewsets.GenericViewSet):
    serializer_class = FileHistorySerializer
    def get_queryset(self,pk=None):
        return File.historical.filter(id = pk,history_user_id=self.userFull.id)
    
    def retrieve(self,request,pk = None):
        fileResult = File.objects.get(slug = pk,unidadArea_id = self.userFull.unidadArea_id)
        if fileResult:
            historyFile = self.get_queryset(fileResult.id)
            if validarPrivado(fileResult,self.userFull.id,True):
                return Response({'error':'La carpeta contenedora o el file son privados'},status = status.HTTP_401_UNAUTHORIZED)
            fileHistorialSerializer = self.get_serializer(historyFile,many = True)
            return Response(fileHistorialSerializer.data, status = status.HTTP_200_OK)
        return Response({'error':'No existe el file'},status = status.HTTP_401_UNAUTHORIZED)
class FileDeleteAPIView(Authentication,viewsets.GenericViewSet):
    serializer_class = FileDetalleSerializer
    def get_queryset(self,pk):
        return self.get_serializer().Meta.model.objects.filter(slug = pk,scope = False,user_id = self.userFull.id,unidadArea_id = self.userFull.unidadArea_id)    
    def retrieve(self,request,pk):
        fileResult = self.get_queryset(pk).first()
        if fileResult:
            #fileSerializer = self.get_serializer(fileResult)
            if fileResult.eliminado:
                rutaFile = settings.MEDIA_ROOT+'files/'
                os.remove(os.path.join(rutaFile+fileResult.documento_file.name))
                fileResult.delete()
                return Response({'mensaje':'File destruido correctamente'}, status = status.HTTP_200_OK)
            folderMasterPrivate = Folder.objects.get(scope = False,
                                                        unidadArea_id = self.userFull.unidadArea_id,
                                                        user_id =self.userFull.id,
                                                        carpeta_hija__isnull =True,eliminado = False)
            fileResult.eliminado = True
            fileResult.save()
            FileInFolder.objects.filter(file_id = fileResult.id).update(parent_folder_id = folderMasterPrivate.id)
            return Response({'mensaje':'File eliminado correctamente'}, status = status.HTTP_200_OK)
        return Response({'error':'No existe el file'},status = status.HTTP_401_UNAUTHORIZED)
class FileRestaurarAPIView(Authentication,viewsets.GenericViewSet):
    serializer_class = FileDetalleSerializer
    def get_queryset(self,pk):
        return self.get_serializer().Meta.model.objects.filter(slug = pk,scope = False,user_id = self.userFull.id,unidadArea_id = self.userFull.unidadArea_id,eliminado = True)    
    def retrieve(self,request,pk):
        fileResult = self.get_queryset(pk).first()
        if fileResult:
            folderMasterPrivate = Folder.objects.get(scope = False,
                                                        unidadArea_id = self.userFull.unidadArea_id,
                                                        user_id =self.userFull.id,
                                                        carpeta_hija__isnull =True,eliminado = False)
            fileResult.eliminado = False
            fileResult.save()
            FileInFolder.objects.filter(file_id = fileResult.id).update(parent_folder_id = folderMasterPrivate.id)
            return Response({'mensaje':'File restaurado a la carpeta privada correctamente'}, status = status.HTTP_200_OK)
        return Response({'error':'No existe el file'},status = status.HTTP_401_UNAUTHORIZED)

def obtenerTextoPDF(file):
        text =""
        
        '''with pdfplumber.open(settings.MEDIA_ROOT+'files'+file) as pdf:
            pdf.strict = True
            totalpages = len(pdf.pages)
            for x in range(0,totalpages):
                first_page = pdf.pages[x]
                text = text + str(first_page.extract_text().replace("",''))
                percent = ((x+1)*100)/totalpages
                print(str(round(percent,2))+" %")
                #print(first_page.extract_text().replace("",''))
        return text'''
        #documentoRenderisado = DocumentoOCR(file)
        try:

            pdfFileObj = open(settings.MEDIA_ROOT +'files'+file,'rb')
            
            pdfReader = PyPDF2.PdfFileReader(pdfFileObj,  strict = False)
            
            #print(pdfReader.isEncrypted)
            #print(pdfReader.numPages)
            #print(pdfReader.getPage(0).extractText())
            for x in range(0, pdfReader.numPages):
                pageObj = pdfReader.getPage(x)
                #print(pageObj.extractText())
                text = text + str(pageObj.extractText())
                percent = ((x+1)*100)/pdfReader.numPages
                print(str(round(percent,2))+" %")
            return text 
        except:
            return None


        '''text = extrarText(file)
        return text'''
        #text = str(documentoRenderisado.obtenerTexto())
